"""Render the HTML engineering report into a Word document.

Walks the published HTML with BeautifulSoup and maps each construct onto a
Word equivalent: sections become Heading 1, the collapsible function blocks
become shaded signature blocks, and the figure SVGs are swapped for the PNGs
produced by render_svgs.py.
"""
import os
import re

from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SRC = "arinc615a-cli-engineering.html"
FIGS = "figs"
OUT = "ARINC-615A-CLI-Engineering-Trace.docx"

INK = RGBColor(0x16, 0x1A, 0x20)
MUTED = RGBColor(0x5A, 0x64, 0x72)
ACCENT = RGBColor(0xA8, 0x5D, 0x00)
TEAL = RGBColor(0x0C, 0x6B, 0x69)
DANGER = RGBColor(0xA6, 0x30, 0x1F)

SANS = "Segoe UI"
SANS_BOLD = "Segoe UI Semibold"
SERIF = "Cambria"
MONO = "Consolas"

SHADE_CODE = "F2F3F5"
SHADE_FN = "EFF1F4"
SHADE_NOTE = "EAF2F1"
SHADE_WARN = "FAEDEA"
SHADE_HDR = "E4E7EC"


# ---------------------------------------------------------------- low level

def shade(element, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def para_shade(paragraph, fill):
    shade(paragraph._p.get_or_add_pPr(), fill)


def para_border(paragraph, colour="A85D00", size=18, edges=("left",)):
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    for edge in edges:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "6")
        el.set(qn("w:color"), colour)
        borders.append(el)
    pPr.append(borders)


def keep_with_next(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    el = OxmlElement("w:keepNext")
    pPr.append(el)


def cell_shade(cell, fill):
    shade(cell._tc.get_or_add_tcPr(), fill)


def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def field(paragraph, instruction):
    """Insert a Word field code (used for the TOC and page numbers)."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, sep, end):
        run._r.append(el)
    return run


# ---------------------------------------------------------------- styles

def build_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = SERIF
    normal.font.size = Pt(10)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    def heading(name, size, colour, space_before, bold_font=SANS_BOLD):
        st = styles[name]
        st.font.name = bold_font
        st.font.size = Pt(size)
        st.font.color.rgb = colour
        st.font.bold = True
        st.paragraph_format.space_before = Pt(space_before)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.keep_with_next = True
        return st

    heading("Heading 1", 18, INK, 20)
    heading("Heading 2", 13.5, INK, 14)
    heading("Heading 3", 11.5, RGBColor(0x39, 0x42, 0x4F), 10)
    heading("Title", 30, INK, 0)

    for name, base, font, size, colour in [
        ("Eyebrow", "Normal", SANS, 8, ACCENT),
        ("CodeBlock", "Normal", MONO, 8.5, INK),
        ("FnSig", "Normal", MONO, 9, INK),
        ("SrcRef", "Normal", MONO, 7.5, MUTED),
        ("FigCaption", "Normal", SERIF, 8.5, MUTED),
        ("Lede", "Normal", SERIF, 11.5, RGBColor(0x39, 0x42, 0x4F)),
        ("NoteLabel", "Normal", SANS, 8, TEAL),
    ]:
        st = styles.add_style(name, 1)
        st.base_style = styles[base]
        st.font.name = font
        st.font.size = Pt(size)
        st.font.color.rgb = colour
        st.paragraph_format.space_after = Pt(3)

    styles["Eyebrow"].font.bold = True
    styles["Eyebrow"].paragraph_format.space_before = Pt(16)
    styles["Eyebrow"].paragraph_format.space_after = Pt(0)
    styles["Eyebrow"].paragraph_format.keep_with_next = True
    styles["NoteLabel"].font.bold = True
    styles["FigCaption"].font.italic = True
    styles["CodeBlock"].paragraph_format.space_after = Pt(0)
    styles["CodeBlock"].paragraph_format.line_spacing = 1.0
    styles["FnSig"].font.bold = True

    mono_char = styles.add_style("MonoInline", 2)
    mono_char.font.name = MONO
    mono_char.font.size = Pt(9)
    mono_char.font.color.rgb = RGBColor(0x39, 0x42, 0x4F)
    return styles


# ---------------------------------------------------------------- inline

def add_inline_node(paragraph, child, bold=False, italic=False, mono=False):
    """Add one node - string or inline tag - as runs on the paragraph."""
    if isinstance(child, NavigableString):
        text = re.sub(r"\s+", " ", str(child))
        if not text:
            return
        run = paragraph.add_run(text)
        run.font.bold = bold or None
        run.font.italic = italic or None
        if mono:
            run.font.name = MONO
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x39, 0x42, 0x4F)
        return
    name = child.name
    if name == "code":
        add_inline(paragraph, child, bold, italic, True)
    elif name in ("strong", "b"):
        add_inline(paragraph, child, True, italic, mono)
    elif name in ("em", "i"):
        add_inline(paragraph, child, bold, True, mono)
    elif name == "br":
        paragraph.add_run().add_break()
    else:
        add_inline(paragraph, child, bold, italic, mono)


def add_inline(paragraph, node, bold=False, italic=False, mono=False):
    """Recursively add a node's children as runs, preserving inline formatting."""
    for child in node.children:
        add_inline_node(paragraph, child, bold, italic, mono)


def text_of(node):
    return re.sub(r"\s+", " ", node.get_text(" ", strip=True))


# ---------------------------------------------------------------- blocks

def add_code_block(doc, text, indent=Cm(0)):
    lines = text.rstrip("\n").split("\n")
    for i, line in enumerate(lines):
        p = doc.add_paragraph(style="CodeBlock")
        p.add_run(line if line.strip() else " ")
        para_shade(p, SHADE_CODE)
        p.paragraph_format.left_indent = indent
        if i == 0:
            p.paragraph_format.space_before = Pt(6)
        if i == len(lines) - 1:
            p.paragraph_format.space_after = Pt(10)
    return p


def add_note(doc, node):
    warn = "warn" in node.get("class", [])
    label = node.find("span", class_="lbl")
    fill = SHADE_WARN if warn else SHADE_NOTE
    colour = "A6301F" if warn else "0C6B69"

    if label:
        p = doc.add_paragraph(style="NoteLabel")
        run = p.add_run(text_of(label).upper())
        run.font.color.rgb = DANGER if warn else TEAL
        para_shade(p, fill)
        para_border(p, colour)
        p.paragraph_format.space_before = Pt(8)
        keep_with_next(p)

    paras = node.find_all("p", recursive=False)
    for i, para in enumerate(paras):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8) if i == len(paras) - 1 else Pt(3)
        add_inline(p, para)
        para_shade(p, fill)
        para_border(p, colour)


BLOCK_TAGS = ("pre", "ul", "ol", "div", "figure", "table", "details", "p")


def add_list(doc, node, ordered):
    """Emit a list. A step may carry block children (a command, a nested list);
    those become their own indented blocks rather than being flattened inline."""
    items = node.find_all("li", recursive=False)
    for n, li in enumerate(items, start=1):
        # Walk children in document order so prose that follows a command
        # block stays after it rather than being hoisted into the lead line.
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.9)
        p.paragraph_format.first_line_indent = Cm(-0.9)
        p.paragraph_format.space_after = Pt(3)
        marker = p.add_run(f"{n:02d}   " if ordered else "•   ")
        marker.font.name = MONO
        marker.font.size = Pt(8.5)
        marker.font.color.rgb = ACCENT if ordered else MUTED
        marker.font.bold = ordered

        current = p
        for child in li.children:
            if getattr(child, "name", None) in BLOCK_TAGS:
                if current is not None and current.runs:
                    keep_with_next(current)
                if child.name == "pre":
                    add_code_block(doc, child.get_text(), indent=Cm(0.9))
                elif child.name == "p":
                    bp = doc.add_paragraph()
                    bp.paragraph_format.left_indent = Cm(0.9)
                    bp.paragraph_format.space_after = Pt(3)
                    add_inline(bp, child)
                else:
                    emit(doc, child)
                current = None
            else:
                if isinstance(child, NavigableString) and not child.strip():
                    continue
                if current is None:
                    current = doc.add_paragraph()
                    current.paragraph_format.left_indent = Cm(0.9)
                    current.paragraph_format.space_after = Pt(3)
                add_inline_node(current, child)


def add_table(doc, wrapper):
    table_node = wrapper.find("table")
    head_cells = table_node.find("thead").find_all("th")
    body_rows = table_node.find("tbody").find_all("tr")

    table = doc.add_table(rows=1, cols=len(head_cells))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    hdr = table.rows[0]
    set_repeat_header(hdr)
    for cell, th in zip(hdr.cells, head_cells):
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run(text_of(th).upper())
        run.font.name = SANS_BOLD
        run.font.size = Pt(7.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x39, 0x42, 0x4F)
        cell_shade(cell, SHADE_HDR)

    for tr in body_rows:
        cells = tr.find_all(["td", "th"])
        row = table.add_row()
        for cell, td in zip(row.cells, cells):
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            classes = td.get("class", [])
            if "m" in classes:
                # whole-cell monospace column
                for i, part in enumerate(str(td.decode_contents()).split("<br/>")):
                    frag = BeautifulSoup(part, "html.parser")
                    if i:
                        p = cell.add_paragraph()
                        p.paragraph_format.space_after = Pt(2)
                    run_p = p
                    add_inline(run_p, frag)
                    for run in run_p.runs:
                        run.font.name = MONO
                        run.font.size = Pt(8)
            else:
                add_inline(p, td)
                for run in p.runs:
                    if run.font.size is None:
                        run.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_function_block(doc, node):
    summary = node.find("summary")
    sig = summary.find("span", class_="fn-sig")
    src = summary.find("span", class_="src")

    p = doc.add_paragraph(style="FnSig")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, sig)
    for run in p.runs:
        run.font.name = MONO
        run.font.size = Pt(9)
        run.font.bold = True
    para_shade(p, SHADE_FN)
    para_border(p, "A85D00", 18)
    keep_with_next(p)

    if src:
        sp = doc.add_paragraph(style="SrcRef")
        sp.paragraph_format.space_after = Pt(4)
        sp.add_run(text_of(src))
        para_shade(sp, SHADE_FN)
        para_border(sp, "A85D00", 18)
        keep_with_next(sp)

    body = node.find("div", class_="fn-body")
    if body:
        walk_children(doc, body)


def add_figure(doc, node, fig_paths):
    if not fig_paths:
        return
    path = fig_paths.pop(0)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        p.add_run().add_picture(path, width=Cm(16.0))
    cap = node.find("figcaption")
    if cap:
        c = doc.add_paragraph(style="FigCaption")
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(12)
        add_inline(c, cap)
        for run in c.runs:
            run.font.italic = True


# ---------------------------------------------------------------- walker

FIG_PATHS = []


def walk_children(doc, container):
    for node in container.children:
        if isinstance(node, NavigableString):
            continue
        emit(doc, node)


def emit(doc, node):
    name = node.name
    classes = node.get("class", []) if hasattr(node, "get") else []

    if name == "h2":
        doc.add_heading(text_of(node), level=1)
    elif name == "h3":
        doc.add_heading(text_of(node), level=2)
    elif name == "h4":
        doc.add_heading(text_of(node), level=3)
    elif name == "span" and "sec-no" in classes:
        p = doc.add_paragraph(style="Eyebrow")
        p.add_run(text_of(node).upper())
    elif name == "p":
        style = "Lede" if "lede" in classes else None
        p = doc.add_paragraph(style=style)
        add_inline(p, node)
    elif name == "ol":
        add_list(doc, node, ordered=True)
    elif name == "ul":
        add_list(doc, node, ordered=False)
    elif name == "pre":
        add_code_block(doc, node.get_text())
    elif name == "div" and "note" in classes:
        add_note(doc, node)
    elif name == "div" and "tw" in classes:
        add_table(doc, node)
    elif name == "details" and "fn" in classes:
        add_function_block(doc, node)
    elif name == "figure":
        add_figure(doc, node, FIG_PATHS)
    elif name in ("div", "section", "figure"):
        walk_children(doc, node)


# ---------------------------------------------------------------- document

def page_setup(doc):
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)


def add_running_head(doc, title):
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0]
        hp.text = ""
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run(title)
        run.font.name = SANS
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        field(fp, "PAGE")
        for run in fp.runs:
            run.font.name = MONO
            run.font.size = Pt(8)
            run.font.color.rgb = MUTED


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph(style="Eyebrow")
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("ARINC 615A-4  ·  HOST DATA LOADER")
    run.font.size = Pt(9)

    t = doc.add_paragraph(style="Title")
    t.paragraph_format.space_after = Pt(4)
    t.add_run("CLI Engineering Trace")

    s = doc.add_paragraph(style="Lede")
    s.paragraph_format.space_after = Pt(20)
    s.add_run(
        "A function-by-function walkthrough of the command-line data loader: "
        "every function on the path from main() to a byte on the wire, and back "
        "up through the handler callbacks that print to the console."
    )

    meta = [
        ("Subject", "arinc_615a — ARINC 615A Tool Suite, CLI targets only"),
        ("Covers", "arinc_615a_operation · arinc_615a_download_request_file"),
        ("Sections", "23, including customisation, build and run procedures"),
        ("Method", "Static source reading; no build or live target was exercised"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    for label, value in meta:
        row = table.add_row()
        row.cells[0].width = Cm(3.0)
        row.cells[1].width = Cm(13.0)
        lp = row.cells[0].paragraphs[0]
        lr = lp.add_run(label.upper())
        lr.font.name = SANS_BOLD
        lr.font.size = Pt(7.5)
        lr.font.color.rgb = MUTED
        lp.paragraph_format.space_after = Pt(3)
        vp = row.cells[1].paragraphs[0]
        vr = vp.add_run(value)
        vr.font.size = Pt(9.5)
        vp.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()
    warn = doc.add_paragraph()
    warn.paragraph_format.space_before = Pt(10)
    wr = warn.add_run(
        "Line numbers refer to the working copy this document was generated from. "
        "Verify against your checkout before relying on them."
    )
    wr.font.size = Pt(8.5)
    wr.font.color.rgb = MUTED
    wr.font.italic = True

    doc.add_page_break()


def add_toc(doc):
    h = doc.add_heading("Contents", level=1)
    h.paragraph_format.space_before = Pt(0)

    note = doc.add_paragraph()
    nr = note.add_run(
        "The table below is a Word field. Open the document, select it and press "
        "F9 (or Ctrl+A then F9) to populate page numbers."
    )
    nr.font.size = Pt(8.5)
    nr.font.italic = True
    nr.font.color.rgb = MUTED

    p = doc.add_paragraph()
    field(p, r'TOC \o "1-2" \h \z \u')
    doc.add_page_break()


def main():
    soup = BeautifulSoup(open(SRC, encoding="utf-8").read(), "html.parser")

    for idx in range(1, 99):
        path = os.path.join(FIGS, f"fig{idx}.png")
        if os.path.exists(path):
            FIG_PATHS.append(path)
        else:
            break

    total_figs = len(FIG_PATHS)

    doc = Document()
    build_styles(doc)
    page_setup(doc)
    add_running_head(doc, "ARINC 615A CLI Engineering Trace")
    add_cover(doc)
    add_toc(doc)

    sections = soup.find_all("section")
    for i, section in enumerate(sections):
        walk_children(doc, section)
        if i < len(sections) - 1:
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    doc.save(OUT)

    size_kb = round(os.path.getsize(OUT) / 1024, 1)
    print(f"wrote {OUT}  ({size_kb} KB)")
    print(f"  sections: {len(sections)}")
    print(f"  figures embedded: {total_figs - len(FIG_PATHS)} of {total_figs}")
    print(f"  paragraphs: {len(doc.paragraphs)}   tables: {len(doc.tables)}")


if __name__ == "__main__":
    main()
