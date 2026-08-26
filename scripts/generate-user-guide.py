#!/usr/bin/env python3
"""Generate docs/ARINC615A-CLI-Installation-and-Test-Procedure.docx.

The .docx is a generated artifact; this script is its source. Regenerate with:

    python scripts/generate-user-guide.py

Requires: pip install python-docx
"""
import pathlib
import sys

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:  # pragma: no cover
    sys.exit("python-docx is required:  pip install python-docx")

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parent))
import _diagrams  # noqa: E402

REPO = pathlib.Path(__file__).resolve().parent.parent
OUT = REPO / "docs" / "ARINC615A-CLI-Installation-and-Test-Procedure.docx"
FIGDIR = REPO / "docs" / "figures"

ACCENT = RGBColor(0x0B, 0x5C, 0xA8)   # ARINC-ish blue
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5A, 0x64, 0x72)
CODE_BG = "F2F4F7"
OK_BG = "E7F5EC"
WARN_BG = "FDF3E2"


# ----------------------------------------------------------------- helpers --
def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def code(doc, text, bg=CODE_BG):
    """A shaded single-cell table holding monospace text."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = t.cell(0, 0)
    shade(c, bg)
    c.text = ""
    for i, line in enumerate(text.strip("\n").split("\n")):
        p = c.paragraphs[0] if i == 0 else c.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def para(doc, text, size=10, bold=False, color=INK, space=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(it)
        r.font.size = Pt(10)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, "E8EDF3")
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def h(doc, text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs:
        r.font.color.rgb = ACCENT if level <= 2 else MUTED
    return hd


_FIGNO = {"n": 0}


def figure(doc, path, caption, width=6.6):
    """Embed a rendered diagram with a numbered caption."""
    _FIGNO["n"] += 1
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Figure {}  —  {}".format(_FIGNO["n"], caption))
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = MUTED


def testcase(doc, tid, title, objective, pre, steps, expected, status):
    h(doc, "{}  {}".format(tid, title), 3)
    table(doc,
          ["Field", "Detail"],
          [["Objective", objective],
           ["Preconditions", pre],
           ["Verification status", status]],
          widths=[1.5, 5.0])
    para(doc, "Procedure", bold=True, size=10, space=2)
    code(doc, steps)
    para(doc, "Expected result", bold=True, size=10, space=2)
    code(doc, expected, bg=OK_BG)


# -------------------------------------------------------------------- body --
def build():
    figs = _diagrams.render_all(FIGDIR)
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10)

    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.9)
        s.top_margin = s.bottom_margin = Inches(0.8)

    # --- title page ---
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ARINC 615A Tool Suite — CLI")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = ACCENT

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Installation, Operation and Test Procedure")
    r2.font.size = Pt(15)
    r2.font.color.rgb = MUTED

    doc.add_paragraph()
    table(doc, ["Item", "Value"], [
        ["Document title", "ARINC 615A CLI — Installation, Operation and Test Procedure"],
        ["Applies to", "arinc_615a_operation (CLI), all 16 commands"],
        ["Repository", "arinc-615a-tool-suite"],
        ["Platforms", "Windows x64 (MSVC 2022); Linux x64/arm64 (GCC 13+)"],
        ["Document status", "Issue 1"],
        ["Licence", "MPL-2.0 — upstream project by Thomas Vogt"],
    ], widths=[1.8, 4.7])

    para(doc, "Verification note", bold=True, space=2)
    para(doc,
         "Every expected result in Section 9 marked VERIFIED was produced by executing the "
         "stated command on a Windows x64 machine and copying the actual output. Results marked "
         "NOT EXECUTED are specified from the tool's own option definitions but were not run, "
         "because they require either ARINC 615A target hardware on the network or a Linux "
         "environment. No expected output in this document is invented.",
         size=9, color=MUTED)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # --- 1 ---
    h(doc, "1  Purpose and scope", 1)
    para(doc,
         "This document is a step-by-step procedure for installing, operating and testing the "
         "ARINC 615A command-line data loader. It is written to be followed literally: every step "
         "is a command to type, with the output you should see.")
    para(doc, "In scope:", bold=True, space=2)
    bullets(doc, [
        "Installing all prerequisites and dependencies on Windows and Linux.",
        "Building the application from source.",
        "Running each of the 16 commands.",
        "A test procedure with pass/fail criteria.",
        "Diagnosing the failures that are known to occur.",
    ])
    para(doc, "Out of scope:", bold=True, space=2)
    bullets(doc, [
        "The target-side implementation (the avionics box). This tool is the host.",
        "Airworthiness, certification or approval of any data-loading activity.",
        "The ARINC standards themselves, which are published by SAE ITC.",
    ])

    h(doc, "1.1  Acronyms", 2)
    table(doc, ["Term", "Meaning"], [
        ["DLP", "Data Loader Protocol"],
        ["FIND", "Find Identification of Network Devices — the discovery protocol"],
        ["IRQ / IAN", "FIND Information Request / Information Answer"],
        ["LRU", "Line Replaceable Unit — an avionics box"],
        ["MSM", "Media Set Manager — the local ARINC 665 software store"],
        ["THA", "Target Hardware Application — the software on the target that speaks 615A"],
        ["TFTP", "Trivial File Transfer Protocol (RFC 1350), the transport"],
        ["P/N", "Part Number"],
    ], widths=[1.2, 5.3])

    # --- 2 ---
    h(doc, "2  System overview", 1)
    para(doc,
         "ARINC 615A is a file-driven protocol layered on TFTP. The host (this tool) and the "
         "target exchange files with fixed three-letter extensions; the state machine is driven by "
         "which file arrives next. FIND is the exception — a plain UDP request/answer pair used to "
         "discover targets before a transfer.")
    table(doc, ["Extension", "File", "Direction"], [
        ["LCI", "Load Configuration Initialisation", "host to target"],
        ["LCL", "Load Configuration List (part numbers, THW IDs)", "target to host"],
        ["LCS", "Load Configuration Status", "target to host"],
        ["LUI / LUR / LUS", "Upload Initialisation / Request / Status", "both"],
        ["LNR / LNA / LNS", "Download Request / Answer / Status", "both"],
    ], widths=[1.3, 3.6, 1.6])
    figure(doc, figs["system_context"],
           "System context: the tool is the host; the data path is files over TFTP")
    figure(doc, figs["layer_stack"],
           "Six layers, one direction. Calls descend; results return through handlers")
    para(doc,
         "Two behaviours matter operationally. The exception timer (--dlp-timeout, default 13 s) "
         "fails an operation if the target stops sending status files, rather than hanging. And "
         "the first Ctrl-C is a graceful protocol abort, not a kill — it tells the target to stop "
         "so it is not left mid-load in an undefined state. A second Ctrl-C terminates hard.",
         bold=False)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    figure(doc, figs["abort_model"],
           "Abort model: the first Ctrl-C is a protocol abort, the second is a hard stop")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    h(doc, "2.1  Where ARINC 615A sits in the protocol stack", 2)
    para(doc,
         "ARINC 615A is an application-layer protocol. It defines no transport of its own: it "
         "rides on TFTP over UDP over IPv4 over Ethernet. Everything below the session layer is "
         "provided by the operating system and the network, which is why interfacing the loader "
         "to another system is mostly a networking exercise rather than a software one.")
    figure(doc, figs["osi_model"],
           "ARINC 615A mapped onto the OSI reference model, with the ports this build uses")

    h(doc, "2.2  Protocols in use while an operation runs", 2)
    para(doc,
         "A packet capture taken during a load shows more than ARINC 615A. The table lists "
         "everything on the wire and whether this software drives it.")
    table(doc, ["Protocol", "Layer", "Role during an operation", "Driven by"], [
        ["ARINC 615A DLP", "Application",
         "Protocol files LCI, LCL, LCS, LUI, LUR, LUS, LNR, LNA, LNS", "This tool"],
        ["FIND", "Application", "Discovery: IRQ opcode 1, IAN opcode 2", "This tool"],
        ["ARINC 665", "Presentation", "Load and media set format carried in the transfer",
         "This tool"],
        ["TFTP (RFC 1350)", "Session", "Moves every protocol file and every load file",
         "tftp library"],
        ["TFTP options (RFC 2347-2349)", "Session",
         "Block size, timeout and transfer size negotiation", "This tool"],
        ["UDP", "Transport", "Connectionless datagrams; no delivery guarantee",
         "Operating system"],
        ["IPv4", "Network", "Addressing and routing", "Operating system"],
        ["ARP", "Network", "Resolves target IP to a MAC address before any UDP flows",
         "Operating system"],
        ["ICMP", "Network", "Diagnostics only; not part of the protocol", "Operating system"],
        ["Ethernet 802.3", "Data link", "Framing; carries the FIND broadcast", "NIC and switch"],
    ], widths=[1.7, 1.0, 2.6, 1.2])

    h(doc, "2.3  Ports and timers used by this build", 2)
    table(doc, ["Purpose", "Default", "Source constant", "Override"], [
        ["Data load (TFTP)", "UDP 59", "DefaultArinc615aTftpPort", "--server-port"],
        ["FIND discovery", "UDP 1001", "Find::DefaultPort", "--find-port"],
        ["TFTP packet timeout", "2 s", "DefaultArinc615aTftpTimeout", "--tftp-timeout"],
        ["TFTP retries", "1", "DefaultArinc615aTftpRetries", "--dlp-retries"],
        ["DLP exception timer", "13 s", "DefaultArinc615aDlpTimeout", "--dlp-timeout"],
        ["DLP retries", "1", "DefaultArinc615aDlpRetries", "--dlp-retries"],
    ], widths=[1.7, 1.0, 2.3, 1.5])
    para(doc,
         "The data-load port is UDP 59, not the TFTP well-known port 69. ARINC 615A-1 states the "
         "port explicitly and this build takes its default from that constant. If a firewall "
         "between host and target permits only port 69, the load will not start.",
         size=9, color=MUTED)

    figure(doc, figs["protocol_interfaces"],
           "What is exchanged during an operation, and the three integration points")

    h(doc, "2.4  Interfacing the loader to other systems", 2)
    para(doc, "The loader can be driven as a process or as a library; both are supported.")
    bullets(doc, [
        "As a process: invoke the CLI and read its exit code. --targets-list produces "
        "machine-readable discovery output a wrapper can parse and feed back in.",
        "As a library: link lib/arinc_615a and implement the ...OperationHandler interfaces to "
        "receive progress, status and completion callbacks. The CLI is itself only a thin "
        "consumer of that API, so everything it does is available to your own application.",
        "Software to be loaded is supplied as ARINC 665 media sets managed through the Media Set "
        "Manager directory; ARINC 649 supplies the shared check-value functions.",
    ])
    para(doc, "Network conditions that must hold:", bold=True, space=2)
    bullets(doc, [
        "Host and target must share a broadcast domain for FIND to discover anything. Routers "
        "block broadcast; across a router, address the target directly with --target-address.",
        "UDP 59 and UDP 1001 must not be filtered between host and target.",
        "On a multi-homed host set --local-tftp-address and --local-find-address explicitly, or "
        "the sockets may bind to the wrong interface.",
    ])
    figure(doc, figs["deployment"], "Interfacing the loader to an aircraft or bench network")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # --- 3 ---
    h(doc, "3  Prerequisites", 1)
    table(doc, ["Requirement", "Windows", "Linux"], [
        ["Operating system", "Windows 10/11 x64", "x86_64 or aarch64"],
        ["Compiler", "Visual Studio 2022 C++ build tools", "GCC 13 or newer (C++23)"],
        ["Build system", "CMake 4.3+ and Ninja", "CMake 4.3+ and Ninja"],
        ["Source control", "Git", "Git"],
        ["Libraries from", "vcpkg (approx. 84 packages)", "distro packages"],
        ["Privileges", "normal user", "sudo, for package installation"],
        ["Disk space", "approx. 10 GB (vcpkg build trees)", "approx. 2 GB"],
        ["Network", "required — see 3.1", "required — see 3.1"],
    ], widths=[1.5, 2.6, 2.4])

    h(doc, "3.1  Network access required at build time", 2)
    para(doc,
         "The CMake configure step clones five sibling projects with FetchContent. There is no "
         "vendored copy and no offline fallback: without access to git.thomas-vogt.de, configure "
         "cannot complete.")
    table(doc, ["Host", "Needed for"], [
        ["git.thomas-vogt.de", "helper, arinc_649, arinc_665, tftp, commands"],
        ["github.com", "vcpkg clone (Windows); CMake binary (Linux, if distro CMake is too old)"],
        ["Package mirrors", "distro packages (Linux) or vcpkg source downloads (Windows)"],
    ], widths=[2.0, 4.5])

    h(doc, "3.2  Time expectations", 2)
    para(doc,
         "The first Windows build is slow and this is normal, not a fault. vcpkg builds roughly 84 "
         "packages; libiconv alone runs its autotools configure twice and was measured at 2.3 hours "
         "on the reference machine. Subsequent builds are served from the vcpkg binary cache and "
         "take seconds to a few minutes. Linux is minutes, because it uses distro packages and "
         "never builds libiconv from source.")

    # --- 4 ---
    h(doc, "4  Installation — Windows", 1)

    h(doc, "Step 4.1  Install the C++ toolset", 2)
    para(doc, "Run in an elevated PowerShell or Command Prompt:")
    code(doc, 'winget install --id Microsoft.VisualStudio.2022.BuildTools -e '
              '--override "--quiet --add Microsoft.VisualStudio.Workload.VCTools --includeRecommended"')
    para(doc, "Expected: winget reports successful installation. Skip if Visual Studio 2022 with "
              "the C++ workload is already present.", size=9, color=MUTED)

    h(doc, "Step 4.2  Install CMake and Git", 2)
    code(doc, "winget install --id Kitware.CMake -e\n"
              "winget install --id Git.Git -e")
    para(doc, "CMake must be 4.3 or newer. The CMake bundled inside Visual Studio is 3.31 and is "
              "NOT sufficient; the setup script detects this and searches for a qualifying "
              "installation.", size=9, color=MUTED)

    h(doc, "Step 4.3  Obtain the source", 2)
    code(doc, "git clone <repository-url> arinc-615a-tool-suite\n"
              "cd arinc-615a-tool-suite")

    h(doc, "Step 4.4  Install, build and run — one command", 2)
    para(doc, "This is the entire installation. It installs dependencies, builds, and runs the tool:")
    code(doc, "build.bat")
    para(doc, "To build and immediately run a specific command, pass its arguments through:")
    code(doc, "build.bat -c Find\n"
              "build.bat debug -c Find\n"
              "build.bat --no-run")
    para(doc, "Expected output, abbreviated:", bold=True, size=10, space=2)
    code(doc,
         "==============================================================\n"
         " ARINC 615A Tool Suite\n"
         " install dependencies  ->  build (release)  ->  run\n"
         "==============================================================\n"
         "[1/3] Checking toolchain...\n"
         "  OK: Visual Studio  C:\\Program Files (x86)\\Microsoft Visual Studio\\2022\\BuildTools\n"
         "  OK: CMake        C:\\Program Files\\CMake\\bin\\cmake.exe\n"
         "  OK: Ninja\n"
         "[2/3] Setting up vcpkg at C:\\vcpkg...\n"
         "  OK: vcpkg ready\n"
         "[3/3] Installing dependencies from vcpkg.json...\n"
         "[1/2] Configuring...\n"
         "[2/2] Building all targets...\n"
         "==============================================================\n"
         " BUILD OK\n"
         "==============================================================", bg=OK_BG)

    figure(doc, figs["install_windows"],
           "Windows installation flow, with the two hazards marked")

    h(doc, "Step 4.5  What was installed and where", 2)
    table(doc, ["Path", "Contents", "Why here"], [
        ["C:\\vcpkg", "vcpkg clone", "Short path — see 4.6"],
        ["C:\\vb", "vcpkg build trees", "Short path — see 4.6"],
        ["C:\\vp", "vcpkg packages", "Short path — see 4.6"],
        ["C:\\vi", "installed headers, libs, DLLs", "Short path — see 4.6"],
        ["cmake-build-msvc-static-release\\", "build output and the executable", "In the repository"],
    ], widths=[1.9, 2.6, 2.0])

    figure(doc, figs["install_layout"],
           "Where the installation puts things, and the runtime DLL requirement")

    h(doc, "4.6  Why the short paths are mandatory", 2)
    para(doc,
         "vcpkg builds libiconv with autotools. During make install it forms a path by "
         "concatenating the package directory with the full install prefix. Under a deep project "
         "path the result exceeds the Windows 260-character limit, libtool falls back to the "
         "long-path form, and cl.exe rejects it as an unknown option:")
    code(doc,
         "cl : Command line warning D9002 : ignoring unknown option '//?/C:/.../iconv.lib'\n"
         "iconv.obj : error LNK2019: unresolved external symbol libiconv_open\n"
         ".libs\\iconv.exe : fatal error LNK1120: 7 unresolved externals", bg=WARN_BG)
    para(doc,
         "The import library is silently dropped and the port fails after roughly 40 minutes of "
         "work. Do not repoint these roots at longer paths.", size=9, color=MUTED)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # --- 5 ---
    h(doc, "5  Installation — Linux", 1)
    para(doc,
         "Linux uses system libraries rather than vcpkg, because the GCC presets carry no vcpkg "
         "toolchain file. libiconv is therefore never built from source and the failure in 4.6 "
         "cannot occur.")

    h(doc, "Step 5.1  Obtain the source", 2)
    code(doc, "git clone <repository-url> arinc-615a-tool-suite\n"
              "cd arinc-615a-tool-suite")

    h(doc, "Step 5.2  Install, build and run — one command", 2)
    code(doc, "./build.sh\n"
              "./build.sh -c Find\n"
              "./build.sh debug -c Find\n"
              "./build.sh --no-run")
    para(doc, "sudo is invoked for package installation; you will be prompted for your password.",
         size=9, color=MUTED)

    figure(doc, figs["install_linux"], "Linux installation flow: distro packages, no vcpkg")

    h(doc, "Step 5.3  Packages installed per distribution", 2)
    table(doc, ["Need", "apt", "dnf", "pacman", "zypper"], [
        ["Compiler", "build-essential g++", "gcc-c++", "base-devel", "gcc-c++"],
        ["Generator", "ninja-build", "ninja-build", "ninja", "ninja"],
        ["Boost", "libboost-all-dev", "boost-devel", "boost", "boost-devel"],
        ["spdlog", "libspdlog-dev", "spdlog-devel", "spdlog", "spdlog-devel"],
        ["fmt", "libfmt-dev", "fmt-devel", "fmt", "fmt-devel"],
        ["libxml++", "libxml++5.0-dev / 2.6", "libxml++-devel", "libxml++", "libxml++-devel"],
    ], widths=[1.0, 1.6, 1.4, 1.2, 1.4])
    para(doc,
         "If the distribution's CMake is older than 4.3 — most are — the installer downloads the "
         "official Kitware binary into a git-ignored .toolchain/ directory and uses that. Override "
         "the version with CMAKE_VERSION=4.4.2 scripts/install-deps.sh.", size=9, color=MUTED)

    # --- 6 ---
    h(doc, "6  Post-installation verification", 1)
    para(doc, "Three checks confirm a good installation before any operational use.")

    h(doc, "Check 6.1  The executable exists", 2)
    code(doc,
         "REM Windows\n"
         "dir cmake-build-msvc-static-release\\app\\arinc_615a_operation\\arinc_615a_operation.exe\n\n"
         "# Linux\n"
         "ls -l cmake-build-gcc-static-release/app/arinc_615a_operation/arinc_615a_operation")

    h(doc, "Check 6.2  Windows only — put the DLLs on PATH", 2)
    para(doc,
         "The Windows build links against the dynamic x64-windows triplet with "
         "VCPKG_APPLOCAL_DEPS=OFF, so dependency DLLs are not copied next to the executable. "
         "build.bat handles this; running the exe directly does not. Match the configuration — "
         "never mix release DLLs into a debug binary.")
    code(doc,
         'set "PATH=C:\\vi\\x64-windows\\bin;%PATH%"        REM release\n'
         'set "PATH=C:\\vi\\x64-windows\\debug\\bin;%PATH%"  REM debug')
    para(doc, "Linux needs nothing here — system libraries are already on the loader search path.",
         size=9, color=MUTED)

    h(doc, "Check 6.3  The tool responds", 2)
    code(doc, "arinc_615a_operation.exe --help")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # --- 7 ---
    h(doc, "7  Operating procedures", 1)

    h(doc, "7.1  Command syntax", 2)
    code(doc,
         "arinc_615a_operation.exe -c <Command> [options]\n"
         "arinc_615a_operation.exe -c <Command> --help\n"
         "arinc_615a_operation.exe --help")

    h(doc, "7.2  MANDATORY — option syntax", 2)
    para(doc, "Always use the long option form with '='. For example --timeout=5, not -t 5.",
         bold=True)
    para(doc,
         "Space-separated values are mis-consumed, and short flags are reused for different "
         "meanings across option groups, so a short flag can silently bind to the wrong option:")
    table(doc, ["Short flag", "Means", "But also means"], [
        ["-l", "--log-level", "--targets-list"],
        ["-i", "--target-id", "--timeout-option (TFTP)"],
        ["-t", "--timeout (Find)", "--tftp-timeout"],
        ["-p", "--dynamic-port (Find)", "--server-port (TFTP)"],
        ["-d", "--media-set-manager-dir", "--dally (TFTP)"],
        ["-c", "--command", "--check-file-integrity, --no-check-validity"],
    ], widths=[1.2, 2.4, 2.9])

    h(doc, "7.3  Addressing a target", 2)
    code(doc,
         "REM directly by IP\n"
         "--target-address=10.0.0.5 --target-id=THW-1\n\n"
         "REM or resolve the ID through a list produced by Find\n"
         "--targets-list=targets.json --target-id=THW-1")
    para(doc, "--target-id is required on every operation. BatchUpload accepts only "
              "--targets-list, because a batch spans multiple targets.", size=9, color=MUTED)

    h(doc, "7.4  Options common to all ARINC 615A operations", 2)
    table(doc, ["Option", "Default", "Purpose"], [
        ["--dlp-timeout=<s>", "13", "Exception timer; fails a stalled operation"],
        ["--dlp-retries=<n>", "—", "Protocol retry count"],
        ["--port-option", "off", "Enable the ARINC 615A-3 Port Option"],
        ["--local-tftp-address=<ip>", "—", "Interface the TFTP server binds to"],
        ["--log-protocol-files", "off", "Log every protocol file sent and received"],
        ["--log-level=<lvl>", "warn", "trace, debug, info, warn, err, critical, off"],
        ["--server-port=<port>", "—", "Local TFTP server UDP port"],
        ["--tftp-timeout=<s>", "—", "Packet timeout when none is negotiated"],
        ["--dally", "—", "Wait after final ACK to survive a lost last ACK"],
        ["--block-size-option[=n]", "1468", "Negotiate TFTP block size, 8 to 65464"],
        ["--timeout-option[=s]", "2", "Negotiate TFTP timeout, 1 to 255"],
        ["--handle-transfer-size-option", "—", "Negotiate the transfer size option"],
    ], widths=[2.2, 0.9, 3.4])

    h(doc, "7.5  Typical operational session", 2)
    code(doc,
         "REM 1. discover what is on the network\n"
         "arinc_615a_operation.exe -c Find --timeout=5 --targets-list=targets.json\n\n"
         "REM 2. read what is currently loaded on a target\n"
         "arinc_615a_operation.exe -c Information --targets-list=targets.json --target-id=THW-1\n\n"
         "REM 3. stage the software locally\n"
         "arinc_615a_operation.exe -c Create         --media-set-manager-dir=C:\\msm\n"
         "arinc_615a_operation.exe -c ImportMediaSet --media-set-manager-dir=C:\\msm "
         "--source-directory=D:\\medium1\n"
         "arinc_615a_operation.exe -c ListLoads      --media-set-manager-dir=C:\\msm\n\n"
         "REM 4. upload\n"
         "arinc_615a_operation.exe -c Upload --targets-list=targets.json --target-id=THW-1 "
         "--media-set-manager-dir=C:\\msm --media-set-pn=ABC12-3456-7890 --load-header=LOAD001.LUH\n\n"
         "REM 5. confirm the change took effect\n"
         "arinc_615a_operation.exe -c Information --targets-list=targets.json --target-id=THW-1")

    figure(doc, figs["find_sequence"],
           "FIND discovery: one broadcast, N answers, bounded by --timeout")
    figure(doc, figs["information_sequence"],
           "Information operation: LCI out, LCS status loop, LCL result")
    figure(doc, figs["upload_sequence"],
           "Upload: the host becomes a TFTP server and the target pulls the files")

    h(doc, "7.6  Stopping an operation", 2)
    para(doc,
         "Press Ctrl-C once for a graceful protocol abort — the tool tells the target to stop. "
         "Press Ctrl-C a second time only if the first does not return; that terminates hard and "
         "may leave the target in an undefined state. On real hardware, allow the first abort to "
         "complete.")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # --- 8 ---
    h(doc, "8  Test procedure", 1)
    para(doc,
         "Execute in order. Tests 1 to 12 need no target hardware and form the acceptance set for "
         "an installation. Tests 13 to 16 require hardware or a Linux host.")
    table(doc, ["Legend", "Meaning"], [
        ["VERIFIED", "Executed on Windows x64; the expected output shown is the actual output"],
        ["NOT EXECUTED", "Specified from the tool's option definitions; not run"],
    ], widths=[1.5, 5.0])

    figure(doc, figs["test_flow"], "Test sequence and the acceptance gate")

    h(doc, "8.1  Test environment used for verification", 2)
    table(doc, ["Item", "Value"], [
        ["OS", "Windows 11 x64"],
        ["Compiler", "MSVC 14.44 (Visual Studio 2022 Build Tools)"],
        ["CMake / generator", "4.4.2 / Ninja"],
        ["Configuration", "msvc-static-release"],
        ["Executable size", "2,091,520 bytes"],
        ["Network", "No ARINC 615A target hardware present"],
    ], widths=[1.8, 4.7])

    # --- 9 ---
    h(doc, "9  Test cases", 1)

    testcase(doc, "TC-01", "Toolchain prerequisites are detected",
             "Confirm the installer finds the compiler, CMake 4.3+ and Ninja.",
             "Visual Studio 2022 C++ tools and CMake 4.3+ installed.",
             "scripts\\install-deps.bat",
             "[1/3] Checking toolchain...\n"
             "  OK: Visual Studio  C:\\Program Files (x86)\\Microsoft Visual Studio\\2022\\BuildTools\n"
             "  OK: CMake        C:\\Program Files\\CMake\\bin\\cmake.exe\n"
             "  OK: Ninja\n\n"
             "Exit code: 0",
             "VERIFIED")

    testcase(doc, "TC-02", "One-command build from a clean tree",
             "Confirm dependencies install and all targets build.",
             "TC-01 passed. Network access available.",
             "build.bat --no-run",
             "==== BUILD OK ====\n"
             "Executable: ...\\cmake-build-msvc-static-release\\app\\arinc_615a_operation\\"
             "arinc_615a_operation.exe\n\n"
             "Exit code: 0\n"
             "Note: first run may take hours (libiconv). Later runs, seconds.",
             "VERIFIED")

    testcase(doc, "TC-03", "Build artefact is produced",
             "Confirm the executable exists and is non-trivial in size.",
             "TC-02 passed.",
             "dir cmake-build-msvc-static-release\\app\\arinc_615a_operation\\arinc_615a_operation.exe",
             "One file listed, size approximately 2,091,520 bytes (Release).\n"
             "Debug builds are larger, approximately 9 MB.",
             "VERIFIED")

    testcase(doc, "TC-04", "Command catalogue is complete",
             "Confirm all 16 commands register.",
             "TC-03 passed; DLLs on PATH per 6.2.",
             "arinc_615a_operation.exe --help",
             "ARINC 615A Operation - (<git-hash>)\n"
             "Command Options:\n"
             "  -c [ --command ] command Command to execute\n"
             "  -h [ --help ]            Help\n"
             "Commands are:\n"
             " * AdhocUpload   * BatchUpload   * Create         * Find\n"
             " * ImportMediaSet * ImportMediaSetXml * Information * ListBatches\n"
             " * ListLoads     * ListMediaSets * MedDownload    * OpDownload\n"
             " * RemoveMediaSet * Targets      * Upload         * UploadLoads\n\n"
             "PASS if all 16 names appear.  Exit code: 1 (help with no command)",
             "VERIFIED")

    testcase(doc, "TC-05", "Per-command help dispatches",
             "Confirm a sub-command parses and reports its own options.",
             "TC-04 passed.",
             "arinc_615a_operation.exe -c ListMediaSets --help",
             "List all Media Sets registered with the Media Set Manager.\n\n"
             "List ARINC 665 Media Sets Options:\n"
             "  -d [ --media-set-manager-dir ] Directory\n"
             "                                        ARINC 665 Media Set Manager directory.\n"
             "                                        Required.\n\n"
             "Exit code: 0",
             "VERIFIED")

    testcase(doc, "TC-06", "FIND discovery executes and terminates",
             "Confirm the protocol stack opens a socket, broadcasts, honours its timeout and "
             "shuts down cleanly. This is the primary smoke test.",
             "TC-04 passed. Any network interface up.",
             "arinc_615a_operation.exe -c Find",
             "ARINC 615A Operation - (<git-hash>)\n"
             "ARINC 615A FIND Query\n"
             "ARINC 615A FIND Query finished\n\n"
             "Exit code: 0\n"
             "With no target hardware present, zero targets are reported. That is a PASS —\n"
             "the test proves the stack runs and terminates, not that hardware exists.",
             "VERIFIED")

    testcase(doc, "TC-07", "Media Set Manager can be created",
             "Confirm the ARINC 665 local store initialises.",
             "TC-04 passed.",
             "arinc_615a_operation.exe -c Create --media-set-manager-dir=%TEMP%\\msm-test",
             "Create ARINC 665 Media Set Manager\n"
             "Media Set Manager directory: C:\\Users\\<user>\\AppData\\Local\\Temp\\msm-test\n\n"
             "Exit code: 0\n"
             "A file MediaSetManager.json (6 bytes when empty) is created in that directory.",
             "VERIFIED")

    testcase(doc, "TC-08", "Empty store reports empty",
             "Confirm the list commands read the store and report correctly when it is empty.",
             "TC-07 passed.",
             "arinc_615a_operation.exe -c ListMediaSets --media-set-manager-dir=%TEMP%\\msm-test\n"
             "arinc_615a_operation.exe -c ListLoads     --media-set-manager-dir=%TEMP%\\msm-test\n"
             "arinc_615a_operation.exe -c ListBatches   --media-set-manager-dir=%TEMP%\\msm-test",
             "List ARINC 665 Media Sets\n"
             "*** No media sets within media set manger ***\n\n"
             "List ARINC 665 Loads\n"
             "*** No loads within media set manger ***\n\n"
             "List ARINC 665 Batches\n"
             "*** No batches within media set manger ***\n\n"
             "Exit code: 0 for all three.\n"
             "Note: 'manger' is the tool's own spelling; match it exactly when parsing output.",
             "VERIFIED")

    testcase(doc, "TC-09", "Negative — no command supplied",
             "Confirm a missing command is rejected with guidance.",
             "TC-04 passed.",
             "arinc_615a_operation.exe",
             "Error parsing command line: the option 'command' is required but missing\n"
             "Enter <path>\\arinc_615a_operation.exe --help for command line description.\n\n"
             "Exit code: 1",
             "VERIFIED")

    testcase(doc, "TC-10", "Negative — unknown command",
             "Confirm an invalid command name is rejected.",
             "TC-04 passed.",
             "arinc_615a_operation.exe -c NoSuchCommand",
             "Error parsing command line: the argument ('NoSuchCommand') for option is invalid\n"
             "Enter <path> --command NoSuchCommand --help for command line description.\n\n"
             "Exit code: 1",
             "VERIFIED")

    testcase(doc, "TC-11", "Negative — required option omitted",
             "Confirm required options are enforced per command.",
             "TC-04 passed.",
             "arinc_615a_operation.exe -c Information --target-address=127.0.0.1\n"
             "arinc_615a_operation.exe -c ListMediaSets",
             "ARINC 615A Information Operation\n"
             "Error parsing command line: the option '--target-id' is required but missing\n\n"
             "List ARINC 665 Media Sets\n"
             "Error parsing command line: the option '--media-set-manager-dir' is required but missing\n\n"
             "Exit code: 1 for both",
             "VERIFIED")

    testcase(doc, "TC-12", "Release package is self-contained",
             "Confirm the published binary runs with no vcpkg directory on PATH.",
             "Release zip downloaded and extracted.",
             "REM in a shell whose PATH contains only C:\\Windows and C:\\Windows\\system32\n"
             "arinc_615a_operation.exe -c Find",
             "ARINC 615A Operation - (<git-hash>)\n"
             "ARINC 615A FIND Query\n"
             "ARINC 615A FIND Query finished\n\n"
             "Exit code: 0\n"
             "PASS requires no missing-DLL dialog. The package bundles six DLLs:\n"
             "boost_program_options, fmt, iconv-2, libxml2, xml++, z.",
             "VERIFIED")

    testcase(doc, "TC-13", "Information operation against a target",
             "Read part numbers and versions from real target hardware.",
             "An ARINC 615A target reachable on the network; its IP and target ID known.",
             "arinc_615a_operation.exe -c Information --target-address=<ip> --target-id=<id> "
             "--log-level=info",
             "The host writes an LCI, receives LCS status files while the target works, and\n"
             "finally an LCL listing target hardware IDs and loaded part numbers.\n\n"
             "PASS: part numbers are printed and the exit code is 0.\n"
             "FAIL: operation aborted by DLP (no status within --dlp-timeout, default 13 s).",
             "NOT EXECUTED — requires target hardware")

    testcase(doc, "TC-14", "Upload operation",
             "Upload a load from a registered media set to a target.",
             "TC-13 passed. A media set imported and its P/N and load header known.",
             "arinc_615a_operation.exe -c Upload --target-address=<ip> --target-id=<id> "
             "--media-set-manager-dir=<dir> --media-set-pn=<pn> --load-header=<file>",
             "Progress is reported from LUS status files until a terminating status.\n\n"
             "PASS: exit code 0, and a subsequent TC-13 reports the new part number.\n"
             "Verify the change actually took effect; do not rely on exit code alone.",
             "NOT EXECUTED — requires target hardware")

    testcase(doc, "TC-15", "Download operation",
             "Retrieve files from a target.",
             "TC-13 passed.",
             "REM list what is offered without downloading — omit --file\n"
             "arinc_615a_operation.exe -c MedDownload --target-address=<ip> --target-id=<id>\n\n"
             "REM then fetch specific files\n"
             "arinc_615a_operation.exe -c MedDownload --target-address=<ip> --target-id=<id> "
             "--file=<name> --download-base-directory=<dir>",
             "Without --file the operation cancels after the list transfer, showing what is\n"
             "available. With --file, the named files are written to a sub-directory of\n"
             "--download-base-directory (default: current directory).\n\n"
             "PASS: requested files present on disk and validity check passes; exit code 0.",
             "NOT EXECUTED — requires target hardware")

    testcase(doc, "TC-16", "Linux build and run",
             "Confirm the Linux path installs, builds and runs.",
             "A supported Linux distribution with sudo and network access.",
             "./build.sh --no-run\n"
             "cmake-build-gcc-static-release/app/arinc_615a_operation/arinc_615a_operation -c Find",
             "BUILD OK, then the same FIND output as TC-06, exit code 0.\n"
             "No PATH configuration is needed on Linux.\n\n"
             "The Linux scripts are syntax-checked only. Distro package names are the most\n"
             "likely thing to need adjusting.",
             "NOT EXECUTED — no Linux environment was available")

    h(doc, "9.1  Acceptance criteria", 2)
    para(doc,
         "An installation is acceptable when TC-01 through TC-12 all pass. TC-13 through TC-15 "
         "form the operational acceptance set and can only be executed against target hardware. "
         "TC-16 applies to Linux hosts.")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # --- 10 ---
    h(doc, "10  Troubleshooting", 1)
    table(doc, ["Symptom", "Cause", "Action"], [
        ["LNK2019 / LNK1120 building libiconv; 'ignoring unknown option //?/C:/...'",
         "Path length. vcpkg concatenates the install prefix onto the package directory, "
         "exceeding 260 characters; libtool emits the long-path form and cl.exe drops the .lib.",
         "Ensure vcpkg roots are C:\\vcpkg, C:\\vb, C:\\vp, C:\\vi. Do not relocate to longer paths."],
        ["error C2220 with warning C4127 inside a Boost header",
         "The presets combine /WX with /external:templates-, so a warning in a dependency's "
         "headers becomes an error.",
         "Use scripts\\build.bat, which passes -DCMAKE_COMPILE_WARNING_AS_ERROR=OFF."],
        ["CMake Error: No target \"arinc_615a_download_request_file\"",
         "InstallPackage.cmake referenced a target that app/CMakeLists.txt no longer builds; "
         "fails at generate time, after configure succeeds.",
         "Already fixed here with if( TARGET ... ). Re-apply if merging upstream changes."],
        ["Configure fails cloning dependencies",
         "No network access to git.thomas-vogt.de. There is no offline fallback.",
         "Restore access, or place checkouts beside the source tree so CMake picks them up."],
        ["Application fails to start with a missing-DLL error",
         "Dynamic triplet with VCPKG_APPLOCAL_DEPS=OFF; DLLs are not beside the executable.",
         "Set PATH per 6.2, matching debug/release, or use the Release zip which bundles them."],
        ["CMake reports a version below 4.3",
         "The CMake bundled with Visual Studio is 3.31 and appears first on PATH after vcvars.",
         "Install CMake 4.3+; the script searches for a qualifying installation."],
        ["'the required argument for option --timeout is missing'",
         "Space-separated short option.",
         "Use --timeout=5, not -t 5."],
        ["'the argument (--timeout) for option --log-level is invalid'",
         "-l is bound to both --log-level and --targets-list in Find.",
         "Avoid -l on that sub-command; use long forms with '='."],
        ["Operation aborted by DLP",
         "The target sent no status file within the exception timer.",
         "Check connectivity and target state; raise --dlp-timeout; add --log-protocol-files."],
        ["Build appears frozen for a long time",
         "libiconv's autotools configure. vcpkg buffers a port's output until it finishes.",
         "Normal. Watch file timestamps under C:\\vb\\libiconv rather than the console."],
    ], widths=[1.9, 2.4, 2.2])

    # --- 11 ---
    h(doc, "11  Appendices", 1)
    h(doc, "Appendix A  Exit codes", 2)
    table(doc, ["Code", "Meaning"], [
        ["0", "Success, including a sub-command --help"],
        ["1", "Argument error, protocol failure, or top-level --help with no command"],
    ], widths=[0.8, 5.7])

    h(doc, "Appendix B  Command index", 2)
    table(doc, ["Command", "Group", "Needs target hardware"], [
        ["Find", "Discovery", "No — safe smoke test"],
        ["Targets", "Discovery", "No — reads a saved list"],
        ["Information", "615A operation", "Yes"],
        ["Upload", "615A operation", "Yes"],
        ["AdhocUpload", "615A operation", "Yes"],
        ["BatchUpload", "615A operation", "Yes"],
        ["UploadLoads", "615A operation", "Yes"],
        ["MedDownload", "615A operation", "Yes"],
        ["OpDownload", "615A operation", "Yes"],
        ["Create", "ARINC 665 store", "No"],
        ["ImportMediaSet", "ARINC 665 store", "No"],
        ["ImportMediaSetXml", "ARINC 665 store", "No"],
        ["ListMediaSets", "ARINC 665 store", "No"],
        ["ListLoads", "ARINC 665 store", "No"],
        ["ListBatches", "ARINC 665 store", "No"],
        ["RemoveMediaSet", "ARINC 665 store", "No"],
    ], widths=[1.8, 2.0, 2.7])

    h(doc, "Appendix C  Related documents", 2)
    table(doc, ["Document", "Contents"], [
        ["README.md", "Overview, diagrams, quick start, command list"],
        ["docs/BUILD.md", "Build stages, dependencies, failure modes, linked resources"],
        ["docs/ARCHITECTURE.md", "Layer map, directory walkthrough, design conventions"],
        ["docs/CODE-TRACE.md", "Function-by-function trace from main() to the wire"],
    ], widths=[2.0, 4.5])

    h(doc, "Appendix D  Standards", 2)
    table(doc, ["Reference", "Title"], [
        ["ARINC 615A-4", "Software Data Loader Using Ethernet Interface"],
        ["ARINC 665-5", "Loadable Software Standards"],
        ["ARINC 649", "Common Terminology and Functions for Software Distribution and Loading"],
        ["RFC 1350", "The TFTP Protocol (Revision 2)"],
        ["RFC 2347/2348/2349", "TFTP Option Extension, Blocksize, Timeout and Transfer Size"],
    ], widths=[1.8, 4.7])

    para(doc, "")
    para(doc,
         "ARINC standards are published by SAE ITC and are not redistributed with this software. "
         "ARINC is a trademark of its respective owner; this project implements the publicly "
         "documented protocol and is not affiliated with or endorsed by ARINC.",
         size=8, color=MUTED)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print("wrote", OUT, OUT.stat().st_size, "bytes")


if __name__ == "__main__":
    build()
